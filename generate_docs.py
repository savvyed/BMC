"""
Generates BMC_CMS_Guide.docx — the content editor guide for the
BMC Digital Health Navigator CMS.
Run: python3 generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colors ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1a, 0x52, 0x76)   # --color-primary
TEAL      = RGBColor(0x14, 0x8f, 0x77)   # --color-accent
BLUE      = RGBColor(0x2e, 0x86, 0xc1)   # --color-primary-light
AMBER     = RGBColor(0x78, 0x42, 0x12)   # --color-external
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF4, 0xF6, 0xF7)
BORDER    = RGBColor(0xD5, 0xD8, 0xDC)
MUTED     = RGBColor(0x71, 0x7D, 0x7E)
DARK_TEXT = RGBColor(0x1c, 0x28, 0x33)
WARN_BG   = RGBColor(0xFF, 0xF3, 0xCD)
WARN_BDR  = RGBColor(0xE0, 0xA8, 0x00)
TEAL_BG   = RGBColor(0xD1, 0xF2, 0xEB)
NAVY_BG   = RGBColor(0xD6, 0xE4, 0xF0)


# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom),
                        ('left', left), ('right', right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   'single')
            el.set(qn('w:sz'),    '4')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'),
                   f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, italic=False, size=None,
            color=None, font='Calibri'):
    run = para.add_run(text)
    run.font.name  = font
    run.font.bold  = bold
    run.font.italic = italic
    if size:  run.font.size  = Pt(size)
    if color: run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    para.paragraph_format.space_after  = Pt(6)
    if level == 1:
        add_run(para, text, bold=True, size=18, color=NAVY)
        # Bottom border
        pPr  = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '4')
        bot.set(qn('w:color'), f'{NAVY[0]:02X}{NAVY[1]:02X}{NAVY[2]:02X}')
        pBdr.append(bot)
        pPr.append(pBdr)
    elif level == 2:
        add_run(para, text, bold=True, size=13, color=NAVY)
    elif level == 3:
        add_run(para, text, bold=True, size=11, color=BLUE)
    return para


def add_body(doc, text, size=10.5, color=DARK_TEXT, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after  = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    add_run(para, text, size=size, color=color)
    return para


def add_bullet(doc, text, bold_prefix=None, size=10.5, indent_level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.left_indent  = Inches(0.25 + indent_level * 0.25)
    if bold_prefix:
        add_run(para, bold_prefix + ': ', bold=True, size=size, color=DARK_TEXT)
    add_run(para, text, size=size, color=DARK_TEXT)
    return para


def add_numbered(doc, text, bold_prefix=None, size=10.5):
    para = doc.add_paragraph(style='List Number')
    para.paragraph_format.space_after  = Pt(3)
    para.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        add_run(para, bold_prefix + ' ', bold=True, size=size, color=DARK_TEXT)
    add_run(para, text, size=size, color=DARK_TEXT)
    return para


def add_callout(doc, label, text, bg=NAVY_BG, border_color=NAVY):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_border(cell, top=border_color, bottom=border_color,
                    left=border_color, right=border_color)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(4)
    add_run(p, label + '  ', bold=True, size=10, color=border_color)
    add_run(p, text, size=10, color=DARK_TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_spacer(doc, size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size)


def make_table(doc, headers, rows, col_widths=None):
    """Renders a styled table with a navy header row."""
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    total = Inches(6.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            tbl.columns[i].width = w
    else:
        w = total // n_cols
        for col in tbl.columns:
            col.width = w

    # Header row
    hrow = tbl.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, NAVY)
        set_cell_border(cell, top=NAVY, bottom=NAVY, left=NAVY, right=NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Pt(4)
        add_run(p, hdr, bold=True, size=9.5, color=WHITE)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        bg  = LIGHT_BG if r_idx % 2 == 0 else WHITE
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_border(cell, top=BORDER, bottom=BORDER,
                            left=BORDER, right=BORDER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Pt(4)
            # Support bold prefix with pipe: "**bold** rest"
            if isinstance(cell_text, tuple):
                bold_part, rest = cell_text
                add_run(p, bold_part, bold=True, size=9.5, color=DARK_TEXT)
                add_run(p, rest, size=9.5, color=DARK_TEXT)
            elif cell_text.startswith('⚠️'):
                add_run(p, cell_text, bold=True, size=9.5, color=AMBER)
            else:
                add_run(p, cell_text, size=9.5, color=DARK_TEXT)

    add_spacer(doc, 8)
    return tbl


# ── Document ─────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Cover ────────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.paragraph_format.space_before = Pt(24)
    cover.paragraph_format.space_after  = Pt(4)
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(cover, 'BMC Digital Health Navigator', bold=True, size=22, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(2)
    add_run(sub, 'Content Management System — Editor Guide', bold=False,
            size=14, color=BLUE)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(32)
    add_run(meta, 'April 2026  ·  For BMC and SavvyEd content staff',
            size=9.5, color=MUTED)

    add_spacer(doc, 4)

    # ── 1. Overview ──────────────────────────────────────────────────────────
    add_heading(doc, '1  Overview', level=1)
    add_body(doc, (
        'The BMC Digital Health Navigator site uses Decap CMS — a browser-based '
        'content editor that lets you update text, steps, and scenario content '
        'without touching any code. Changes you save are published to the live site '
        'automatically, usually within 60 seconds.'
    ))
    add_body(doc, 'The CMS manages two types of content:')
    add_bullet(doc, 'Course Challenges — the text that appears in each of the six '
               'learning challenges (Watch, Practice, and Do it steps), in all five '
               'languages.')
    add_bullet(doc, 'How-To Topics — the step-by-step instructions and topic titles '
               'that appear in the Knowledge Base and are also pulled into challenge '
               'Watch steps automatically.')

    add_callout(doc, 'Note',
        'The CMS does not manage videos, H5P interactions, or site layout. '
        'Those are handled separately by a developer.',
        bg=NAVY_BG, border_color=NAVY)

    # ── 2. Accessing the CMS ─────────────────────────────────────────────────
    add_heading(doc, '2  Accessing the CMS', level=1)

    add_heading(doc, '2.1  URL', level=2)
    add_body(doc, 'Open the CMS in any browser at:')
    url_para = doc.add_paragraph()
    url_para.paragraph_format.left_indent  = Inches(0.3)
    url_para.paragraph_format.space_after  = Pt(8)
    add_run(url_para,
            'https://tiny-hamster-9bcc46.netlify.app/admin/',
            bold=True, size=11, color=BLUE)

    add_heading(doc, '2.2  Logging in', level=2)
    add_numbered(doc, 'Go to the URL above.')
    add_numbered(doc, 'Click Log in with Netlify Identity.')
    add_numbered(doc, 'Enter your email address and password. '
                 '(If this is your first time, check your email for an invite '
                 'from Netlify.)')
    add_numbered(doc, 'You will land on the CMS dashboard.')
    add_spacer(doc, 6)

    add_callout(doc, 'First-time setup',
        'If you have not received a login invite, contact your site administrator. '
        'Only accounts approved in Netlify Identity can log in — the CMS is not '
        'accessible to the public.',
        bg=TEAL_BG, border_color=TEAL)

    add_heading(doc, '2.3  Navigating the interface', level=2)
    add_body(doc, 'Once logged in, you will see:')
    add_bullet(doc, 'Left sidebar — the collection list. '
               'Click a collection name to open it.',
               bold_prefix='Sidebar')
    add_bullet(doc, 'Content area — the list of items in the selected collection, '
               'or the editing form for a selected item.',
               bold_prefix='Content area')
    add_bullet(doc, 'Top right — a Publish button (or Save Draft for unpublished changes).',
               bold_prefix='Publish button')
    add_spacer(doc, 4)

    # ── 3. Collections ───────────────────────────────────────────────────────
    add_heading(doc, '3  Collections', level=1)
    add_body(doc, 'The CMS has two collections, shown in the left sidebar:')
    add_spacer(doc, 4)

    make_table(doc,
        headers=['Collection', 'What it contains', 'How often updated'],
        rows=[
            ('Course Challenges',
             'Text for all 6 challenges × 5 languages — headings, scenario text, '
             'checklist items, progress labels, and video captions.',
             'Regularly during content development'),
            ('How-To Topics',
             'Step-by-step instructions for ~28 topics across 7 categories. '
             'Shown in the Knowledge Base and pulled into challenge Watch steps.',
             'Regularly during content development'),
        ],
        col_widths=[Inches(1.6), Inches(3.5), Inches(1.4)],
    )

    # ── 4. Editing challenge content ─────────────────────────────────────────
    add_heading(doc, '4  Editing Challenge Content', level=1)

    add_heading(doc, '4.1  Opening a challenge', level=2)
    add_numbered(doc, 'Click Course Challenges in the left sidebar.')
    add_numbered(doc, 'Click the challenge you want to edit, '
                 'e.g. Challenge 2: Join a video visit.')
    add_numbered(doc, 'The editing form opens. Each challenge has five language '
                 'sections: English, Español, Português, Kreyòl ayisyen, and 中文. '
                 'English is expanded by default; other languages are collapsed.')
    add_numbered(doc, 'Click a language section header to expand it.')
    add_spacer(doc, 4)

    add_heading(doc, '4.2  Field reference — one challenge', level=2)
    add_body(doc, (
        'Every challenge has the same set of fields, organised by step. '
        'The table below describes each field, what it does, and where it '
        'appears on screen.'
    ))
    add_spacer(doc, 4)

    make_table(doc,
        headers=['Field name', 'Where it appears', 'Notes'],
        rows=[
            (('Page title', ''),
             'Browser tab and page header bar',
             'Example: Challenge 02 — Join a video visit'),

            (('Watch — progress label', ''),
             'Small text directly above the main heading on Step 1',
             'Example: Step 1 of 3 — Watch'),

            (('Watch — main heading', ''),
             'Large H1 heading above the video on Step 1',
             'Example: How to join a video visit on your phone'),

            (('Watch — video caption', ''),
             'Text shown on the video placeholder until the real video is added',
             'Briefly describes what the video covers'),

            (('Practice — progress label', ''),
             'Small text above the heading on Step 2',
             'Example: Step 2 of 3 — Practice'),

            (('Practice — main heading', ''),
             'Large heading on Step 2',
             'Names the character and the task they need help with'),

            (('Practice — scenario, sentence 1', ''),
             'First sentence of the character intro on Step 2',
             'Describes the character\'s situation. End with a period.'),

            (('Practice — scenario, sentence 2', ''),
             'Second sentence of the character intro on Step 2',
             'Tells the patient what to help the character do. End with a period.'),

            (('Do it — progress label', ''),
             'Small text above the heading on Step 3',
             'Example: Step 3 of 3 — Do it on your phone'),

            (('Do it — main heading', ''),
             'Large heading on Step 3',
             'Example: Now try it on your own MyChart'),

            (('Do it — intro text', ''),
             'Short sentence shown above the checklist on Step 3',
             'Example: Open MyChart on your phone and follow these steps:'),

            (('Checklist — item 1, item 2, …', ''),
             'Individual checkbox items the patient checks off on Step 3',
             'Start each item with a verb: Tap, Open, Enter, Find...'),
        ],
        col_widths=[Inches(2.1), Inches(2.5), Inches(1.9)],
    )

    add_heading(doc, '4.3  Saving and publishing', level=2)
    add_numbered(doc, 'Make your edits in the form fields.')
    add_numbered(doc,
                 'Click Save draft to save without publishing, or '
                 'Publish to save and push the change to the live site.')
    add_numbered(doc,
                 'After publishing, Netlify automatically rebuilds and redeploys '
                 'the site. This takes approximately 30–60 seconds. Refresh the '
                 'live site to see your changes.')
    add_spacer(doc, 6)

    add_callout(doc, 'Tip',
        'Use Save draft while you are still reviewing or getting translations '
        'approved. Publish only when the content is final.',
        bg=TEAL_BG, border_color=TEAL)

    # ── 5. Editing how-to topics ─────────────────────────────────────────────
    add_heading(doc, '5  Editing How-To Topics', level=1)

    add_heading(doc, '5.1  Opening a topic', level=2)
    add_numbered(doc, 'Click How-To Topics in the left sidebar.')
    add_numbered(doc, 'Click All Topics.')
    add_numbered(doc,
                 'Expand the Topic catalog section, then expand the '
                 'category you want to edit '
                 '(e.g. Video visits).')
    add_numbered(doc,
                 'Find the topic in the Topics in this category list '
                 'and click Edit.')
    add_spacer(doc, 4)

    add_heading(doc, '5.2  Field reference — one topic', level=2)
    add_spacer(doc, 4)

    make_table(doc,
        headers=['Field name', 'Where it appears', 'Notes'],
        rows=[
            (('Title', ''),
             'Topic page heading in the Knowledge Base, and as the label on the '
             'video placeholder in the challenge Watch step',
             'Plain language. Example: Join a video visit on your phone'),

            ('⚠️ URL identifier',
             'Used in page links throughout the site',
             '⚠️ Do not change. Changing this will break links in the Knowledge '
             'Base and in challenge pages.'),

            (('Step-by-step instructions', ''),
             'Numbered list on the topic page and in the challenge Watch step',
             'Each item is one instruction. Start with a verb: Tap, Open, Enter, '
             'Find... Add, remove, or reorder items using the list controls.'),
        ],
        col_widths=[Inches(1.8), Inches(2.6), Inches(2.1)],
    )

    add_heading(doc, '5.3  Adding and removing steps', level=2)
    add_body(doc,
             'Each topic\'s step list has Add step and remove (×) controls in '
             'the CMS. Steps can also be dragged to reorder them.')
    add_bullet(doc, 'Click Add step to add a new instruction at the bottom of the list.')
    add_bullet(doc, 'Click the × icon next to any step to remove it.')
    add_bullet(doc, 'Drag the handle (⠿) on any step to change its order.')
    add_spacer(doc, 4)

    add_callout(doc, 'Remember',
        'How-To Topics are a shared content pool. When you edit a topic here, '
        'the change appears in BOTH the Knowledge Base topic page AND in the '
        'corresponding challenge Watch step automatically. You only need to edit '
        'it once.',
        bg=TEAL_BG, border_color=TEAL)

    # ── 6. Translations ──────────────────────────────────────────────────────
    add_heading(doc, '6  Adding and Updating Translations', level=1)

    add_heading(doc, '6.1  How language fallback works', level=2)
    add_body(doc,
             'If a field in a non-English language section is left blank, the site '
             'automatically falls back to the English text for that field. This means '
             'partial translations are safe — patients will see English for any '
             'untranslated fields rather than a blank or broken page.')

    add_heading(doc, '6.2  Adding a translation', level=2)
    add_numbered(doc, 'Open the challenge or topic you want to translate.')
    add_numbered(doc,
                 'Click the language section header (e.g. Español) to expand it.')
    add_numbered(doc, 'Fill in each field with the translated text. '
                 'Leave a field blank if the translation is not yet ready.')
    add_numbered(doc, 'Click Publish when the translation is approved.')
    add_spacer(doc, 4)

    add_callout(doc, 'Translation scope',
        'The How-To Topics catalog does not currently have per-language fields. '
        'Topic step text is in English only and displayed to all patients. '
        'Full translation of topics is planned for a future phase.',
        bg=NAVY_BG, border_color=NAVY)

    add_heading(doc, '6.3  Supported languages', level=2)
    add_spacer(doc, 4)

    make_table(doc,
        headers=['Language', 'Code', 'Status in prototype'],
        rows=[
            ('English',        'en', 'Complete'),
            ('Español',        'es', 'Partial — to be completed'),
            ('Português',      'pt', 'Placeholder — to be completed'),
            ('Kreyòl ayisyen', 'ht', 'Placeholder — to be completed'),
            ('中文',           'zh', 'Placeholder — to be completed'),
        ],
        col_widths=[Inches(2.0), Inches(1.0), Inches(3.5)],
    )

    # ── 7. What not to change ────────────────────────────────────────────────
    add_heading(doc, '7  What Not to Change', level=1)
    add_body(doc,
             'Some fields in the CMS are technical identifiers that must stay '
             'unchanged. Editing them will break links or cause content to '
             'disappear from the site.')
    add_spacer(doc, 4)

    make_table(doc,
        headers=['Field', 'Where', 'Why it must not change'],
        rows=[
            ('⚠️ URL identifier',
             'Any topic in How-To Topics',
             'Used in URLs throughout the site. Changing it breaks '
             'Knowledge Base and challenge page links.'),
            ('⚠️ Challenge topic refs',
             'Not shown in CMS — developer-managed only',
             'Controls which topics appear in each challenge Watch step. '
             'Contact your developer to make changes.'),
        ],
        col_widths=[Inches(1.6), Inches(2.3), Inches(2.6)],
    )

    add_callout(doc, 'When in doubt',
        'If you are unsure whether a change is safe, save as a draft and check '
        'with your developer before publishing.',
        bg=WARN_BG, border_color=WARN_BDR)

    # ── 8. How publishing works ──────────────────────────────────────────────
    add_heading(doc, '8  How Publishing Works', level=1)
    add_body(doc,
             'When you click Publish in the CMS, the following happens automatically:')
    add_spacer(doc, 4)

    make_table(doc,
        headers=['Step', 'What happens', 'Time'],
        rows=[
            ('1', 'Decap CMS saves your changes as a commit to the GitHub repository.',
             'Instant'),
            ('2', 'Netlify detects the new commit and starts a site rebuild.',
             '~5 seconds'),
            ('3', 'The updated site is built and deployed to Netlify\'s global CDN.',
             '~30–60 seconds'),
            ('4', 'Patients loading the site anywhere in the world see the new content.',
             'Immediately after step 3'),
        ],
        col_widths=[Inches(0.5), Inches(4.7), Inches(1.3)],
    )

    add_body(doc,
             'There is no separate deploy step — publishing in the CMS and going '
             'live are the same action. Use Save draft if you need to review before '
             'making a change public.')

    # ── 9. Full schema reference ─────────────────────────────────────────────
    add_heading(doc, '9  Full Schema Reference', level=1)
    add_body(doc,
             'The table below lists every editable field across both collections, '
             'with the content file it is stored in and the language coverage.')
    add_spacer(doc, 4)

    challenges = [
        ('Challenge 0: Set up your email',       'content/challenge-0.json',  4),
        ('Challenge 1: Set up MyChart',          'content/challenge-01.json', 8),
        ('Challenge 2: Join a video visit',      'content/challenge-02.json', 5),
        ('Challenge 3: Common tasks',            'content/challenge-03.json', 5),
        ('Challenge 4: Family account access',   'content/challenge-04.json', 6),
        ('Challenge 5: Community resources',     'content/challenge-05.json', 7),
    ]

    for ch_name, ch_file, n_checks in challenges:
        add_heading(doc, ch_name, level=2)
        add_body(doc, f'File: {ch_file}  ·  Languages: en, es, pt, ht, zh',
                 size=9, color=MUTED, space_after=4)

        fields_rows = [
            ('Page title',                      'string', 'All 5'),
            ('Watch — progress label',          'string', 'All 5'),
            ('Watch — main heading',            'string', 'All 5'),
            ('Watch — video caption',           'string', 'All 5'),
            ('Practice — progress label',       'string', 'All 5'),
            ('Practice — main heading',         'string', 'All 5'),
            ('Practice — scenario, sentence 1', 'string', 'All 5'),
            ('Practice — scenario, sentence 2', 'string', 'All 5'),
            ('Do it — progress label',          'string', 'All 5'),
            ('Do it — main heading',            'string', 'All 5'),
            ('Do it — intro text',              'string', 'All 5'),
        ]
        for i in range(1, n_checks + 1):
            fields_rows.append((f'Checklist — item {i}', 'text (multi-line)', 'All 5'))

        make_table(doc,
            headers=['Field', 'Widget', 'Languages'],
            rows=fields_rows,
            col_widths=[Inches(2.9), Inches(1.7), Inches(1.9)],
        )

    add_heading(doc, 'How-To Topics', level=2)
    add_body(doc, 'File: content/topics.json  ·  Languages: en only (translations planned)',
             size=9, color=MUTED, space_after=4)

    topic_categories = [
        'Setting up your email',
        'Setting up MyChart',
        'Video visits',
        'Medications and refills',
        'Appointments',
        'Messages and your records',
        'Community resources',
    ]

    topics_rows = []
    for cat in topic_categories:
        topics_rows.append((
            f'{cat} — Title',
            'string',
            'English only'
        ))
        topics_rows.append((
            f'{cat} — URL identifier',
            '⚠️ Do not change',
            '—'
        ))
        topics_rows.append((
            f'{cat} — Step-by-step instructions',
            'list of strings',
            'English only'
        ))

    make_table(doc,
        headers=['Field', 'Widget', 'Languages'],
        rows=topics_rows,
        col_widths=[Inches(2.9), Inches(1.7), Inches(1.9)],
    )

    # ── 10. Getting help ─────────────────────────────────────────────────────
    add_heading(doc, '10  Getting Help', level=1)
    add_body(doc,
             'For questions about content and editorial decisions, contact Tianna Tagami.')
    add_body(doc,
             'For technical issues with the CMS or site, contact your developer.')
    add_spacer(doc, 4)

    make_table(doc,
        headers=['Issue', 'Who to contact'],
        rows=[
            ('CMS login problems',                   'Site administrator / developer'),
            ('Content question — what to write',     'Tianna Tagami, M.Ed.'),
            ('Translation question',                 'Tianna Tagami, M.Ed.'),
            ('Topic or challenge not showing correctly', 'Developer'),
            ('Want to add a new challenge or topic', 'Developer'),
            ('URL identifier accidentally changed',  'Developer — revert immediately'),
        ],
        col_widths=[Inches(3.5), Inches(3.0)],
    )

    add_spacer(doc, 12)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_p,
            'BMC Digital Health Navigator  ·  CMS Editor Guide  ·  April 2026',
            size=8.5, color=MUTED)

    return doc


if __name__ == '__main__':
    doc = build()
    out = 'C:/Users/ttian/BMC/BMC_CMS_Editor_Guide.docx'
    doc.save(out)
    print(f'Saved: {out}')
